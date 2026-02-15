from openai import OpenAI
from pathlib import Path
import json
import re

class ReviewerModule:
    """验收员模块 - 智能代码审计和错误分析（增强版）"""
    
    def __init__(self, config):
        self.config = config
        self.client = OpenAI(
            api_key=config['deepseek_api_key'],
            base_url=config['deepseek_base_url']
        )
    
    def review(self, user_input, pm_analysis, code, execution_result, output_path, workspace_path=None, file_data=None):
        """
        验收代码执行结果（增强版：智能错误分析）
        
        Args:
            user_input: 用户输入
            pm_analysis: PM分析结果
            code: 生成的代码
            execution_result: 执行结果
            output_path: 输出路径
            workspace_path: 工作区路径（用于错误分析）
            file_data: 文件数据（用于错误分析）
        """
        
        # 检查是否执行失败
        if not execution_result['success']:
            # 执行失败，进行深度错误分析
            return self._analyze_error_and_suggest_fix(
                user_input, pm_analysis, code, execution_result, 
                workspace_path, file_data
            )
        
        # 执行成功，进行常规验收
        return self._regular_review(
            user_input, pm_analysis, code, execution_result, output_path
        )
    
    def _analyze_error_and_suggest_fix(self, user_input, pm_analysis, code, execution_result, workspace_path, file_data):
        """
        深度错误分析：审计代码和报错内容，结合用户文件分析错误原因
        """
        
        error_msg = execution_result.get('error', '未知错误')
        stdout = execution_result.get('output', '')
        
        # 分析错误类型
        error_type = self._classify_error(error_msg)
        
        # 获取文件信息（用于分析）
        file_info = self._get_file_info_for_analysis(workspace_path, file_data)
        
        # 提取代码中的关键信息
        code_analysis = self._analyze_code_structure(code)
        
        # 构建深度分析提示词
        prompt = f"""你是一个专业的代码审计专家，负责分析代码执行错误并提供精准的修复方案。

【用户需求】
{user_input}

【用户目标】
{pm_analysis.get('user_goal', user_input)}

【期望效果】
{json.dumps(pm_analysis.get('expected_effect', {}), ensure_ascii=False, indent=2)}

【生成的代码】
```python
{code}
```

【执行错误信息】
错误类型: {error_type}
错误输出: {error_msg}
标准输出: {stdout[:500]}

【代码结构分析】
{json.dumps(code_analysis, ensure_ascii=False, indent=2)}

【工作区文件信息】
{file_info}

【核心任务】
作为代码审计专家，你需要：

1. **错误根因分析**：
   - 分析错误发生的根本原因
   - 识别是文件问题、代码逻辑问题还是环境问题
   - 结合工作区文件信息判断文件是否存在、格式是否正确

2. **代码审计**：
   - 审查代码中的问题点
   - 识别可能导致错误的代码行
   - 分析代码逻辑是否符合用户需求

3. **文件兼容性分析**：
   - 检查代码中的文件路径是否正确
   - 检查文件读取方式是否与实际文件格式匹配
   - 检查是否有文件不存在或格式不兼容的问题

4. **精准修复方案**：
   - 提供具体的代码修改建议
   - 说明需要修改哪些代码行
   - 给出修改后的代码片段

【输出格式】
必须严格按照以下JSON格式输出：

```json
{{
    "approved": false,
    "error_analysis": {{
        "error_type": "错误类型（FileNotFoundError/KeyError/ValueError等）",
        "root_cause": "错误根本原因（2-3句话，要具体）",
        "affected_code_lines": ["受影响的代码行或代码片段"],
        "file_issues": [
            {{
                "file": "文件名",
                "issue": "问题描述",
                "expected": "期望的状态",
                "actual": "实际的状态"
            }}
        ]
    }},
    "fix_strategy": {{
        "approach": "修复策略（简要说明如何修复）",
        "priority": "high/medium/low",
        "estimated_difficulty": "easy/medium/hard"
    }},
    "code_modifications": [
        {{
            "location": "需要修改的位置（函数名或代码行描述）",
            "issue": "当前问题",
            "fix": "修复方案",
            "code_before": "修改前的代码片段",
            "code_after": "修改后的代码片段"
        }}
    ],
    "detailed_suggestions": [
        "具体修改建议1（要详细、可执行）",
        "具体修改建议2",
        "具体修改建议3"
    ],
    "feedback": "总结性反馈（告诉Coder如何修复，2-3句话）"
}}
```

【重要提示】
- 要深入分析错误原因，不要只是重复错误信息
- 结合工作区文件信息判断文件是否存在、格式是否正确
- 提供具体的代码修改方案，包括修改前后的代码对比
- 修改建议要可执行，不要模糊不清
- 如果是文件路径问题，要明确指出正确的路径
- 如果是文件格式问题，要说明如何正确读取
- 只输出JSON，不要添加任何解释文字
"""
        
        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个专业的代码审计专家。你必须深入分析错误原因，结合文件信息提供精准的修复方案。你必须严格按照JSON格式输出，不输出任何其他内容。"
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.2,
                max_tokens=3000
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # 提取JSON
            result_json = self._extract_json(result_text)
            
            # 确保必需字段存在
            if 'error_analysis' not in result_json:
                result_json['error_analysis'] = {}
            if 'fix_strategy' not in result_json:
                result_json['fix_strategy'] = {}
            if 'code_modifications' not in result_json:
                result_json['code_modifications'] = []
            if 'detailed_suggestions' not in result_json:
                result_json['detailed_suggestions'] = []
            if 'feedback' not in result_json:
                result_json['feedback'] = "代码执行失败，需要修复"
            
            # 构建返回结果
            return {
                'success': True,
                'approved': False,
                'feedback': result_json.get('feedback', ''),
                'error_analysis': result_json.get('error_analysis', {}),
                'fix_strategy': result_json.get('fix_strategy', {}),
                'code_modifications': result_json.get('code_modifications', []),
                'suggestions': result_json.get('detailed_suggestions', []),
                'issues': [result_json.get('error_analysis', {}).get('root_cause', '代码执行失败')],
                'raw_output': result_text,
                'needs_deep_fix': True  # 标记需要深度修复
            }
            
        except Exception as e:
            # 异常情况：返回基本错误信息
            return {
                'success': False,
                'approved': False,
                'feedback': f"错误分析失败: {str(e)}",
                'error_analysis': {
                    'error_type': error_type,
                    'root_cause': error_msg
                },
                'suggestions': [f"代码执行错误: {error_msg}"],
                'issues': [error_msg],
                'error': str(e)
            }
    
    def _regular_review(self, user_input, pm_analysis, code, execution_result, output_path):
        """常规验收（代码执行成功时）"""
        
        # 检查输出文件
        output_files = self._get_output_files(output_path)
        
        # 读取输出文件内容（精简版）
        output_content = self._read_output_content(output_path)
        
        # 提取关键信息
        user_goal = pm_analysis.get('user_goal', user_input)
        expected_effect = pm_analysis.get('expected_effect', {})
        success_criteria = pm_analysis.get('success_criteria', [])
        
        # 构建提示词（宽松标准）
        prompt = f"""你是一个务实的质量验收员，负责验收代码执行结果是否基本符合用户需求。

【用户目标】
{user_goal}

【期望效果】
{json.dumps(expected_effect, ensure_ascii=False, indent=2)}

【成功标准】
{json.dumps(success_criteria, ensure_ascii=False, indent=2)}

【代码执行情况】
执行状态: 成功
标准输出: {execution_result.get('output', '无')[:500]}

【输出文件列表】
{output_files}

【输出文件内容】
{output_content}

【验收原则 - 宽松标准】
✅ **通过标准**（满足以下条件即可通过）：
1. 代码成功执行
2. 生成了输出文件
3. 文件格式基本正确
4. 核心数据已处理（70-80%符合即可）
5. 主要功能已实现

【输出格式】
必须严格按照以下JSON格式输出：

```json
{{
    "approved": true/false,
    "has_output": true/false,
    "content_correct": true/false,
    "meets_requirements": true/false,
    "feedback": "验收意见（2-3句话）",
    "issues": ["问题1（如有）"],
    "suggestions": ["改进建议（如有）"]
}}
```

【注意事项】
- 基本符合要求即可通过，不必完美
- 只列出严重问题，忽略小问题
- 采用宽松标准，避免过度迭代
- 只输出JSON，不要添加任何解释文字
"""
        
        try:
            response = self.client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "system",
                        "content": "你是一个务实的质量验收员，采用宽松的验收标准。基本符合要求即可通过。你必须严格按照JSON格式输出，不输出任何其他内容。"
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=1000
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # 提取JSON
            result_json = self._extract_json(result_text)
            
            # 强制规则：没有输出文件则不通过
            if output_files == "输出目录为空":
                result_json['has_output'] = False
                result_json['content_correct'] = False
                result_json['approved'] = False
                if 'issues' not in result_json:
                    result_json['issues'] = []
                result_json['issues'].append("未生成任何输出文件")
            
            # 确保必需字段存在
            if 'feedback' not in result_json:
                result_json['feedback'] = "验收完成"
            if 'issues' not in result_json:
                result_json['issues'] = []
            if 'suggestions' not in result_json:
                result_json['suggestions'] = []
            
            return {
                'success': True,
                'approved': result_json.get('approved', False),
                'feedback': result_json.get('feedback', ''),
                'issues': result_json.get('issues', []),
                'suggestions': result_json.get('suggestions', []),
                'raw_output': result_text
            }
            
        except Exception as e:
            # 异常情况
            return {
                'success': False,
                'approved': False,
                'feedback': f"验收过程出错: {str(e)}",
                'issues': [f"验收异常: {str(e)}"],
                'suggestions': ["请检查代码和执行环境"],
                'error': str(e)
            }
    
    def _classify_error(self, error_msg):
        """分类错误类型"""
        error_msg_lower = error_msg.lower()
        
        if 'filenotfounderror' in error_msg_lower or 'no such file' in error_msg_lower:
            return 'FileNotFoundError'
        elif 'keyerror' in error_msg_lower:
            return 'KeyError'
        elif 'valueerror' in error_msg_lower:
            return 'ValueError'
        elif 'typeerror' in error_msg_lower:
            return 'TypeError'
        elif 'indexerror' in error_msg_lower:
            return 'IndexError'
        elif 'attributeerror' in error_msg_lower:
            return 'AttributeError'
        elif 'importerror' in error_msg_lower or 'modulenotfounderror' in error_msg_lower:
            return 'ImportError'
        elif 'permissionerror' in error_msg_lower:
            return 'PermissionError'
        elif 'syntaxerror' in error_msg_lower:
            return 'SyntaxError'
        else:
            return 'UnknownError'
    
    def _get_file_info_for_analysis(self, workspace_path, file_data):
        """获取文件信息用于错误分析"""
        if not workspace_path or not Path(workspace_path).exists():
            return "工作区信息不可用"
        
        if file_data and file_data.get('files'):
            # 使用已有的文件数据
            files = file_data.get('files', [])
            file_info_lines = ["工作区文件列表:"]
            for f in files:
                file_info_lines.append(
                    f"- {f['name']} (路径: {f['path']}, 类型: {f['extension']}, 大小: {f['size']} bytes)"
                )
            return "\n".join(file_info_lines)
        else:
            # 快速扫描工作区
            workspace = Path(workspace_path)
            files = []
            for file_path in workspace.rglob('*'):
                if file_path.is_file() and not file_path.name.startswith('.'):
                    rel_path = file_path.relative_to(workspace)
                    files.append(f"- {file_path.name} (路径: {rel_path}, 大小: {file_path.stat().st_size} bytes)")
            
            if files:
                return "工作区文件列表:\n" + "\n".join(files)
            else:
                return "工作区为空"
    
    def _analyze_code_structure(self, code):
        """分析代码结构"""
        analysis = {
            'imports': [],
            'functions': [],
            'file_operations': [],
            'variables': []
        }
        
        try:
            # 提取import语句
            import_pattern = r'^(?:from\s+[\w.]+\s+)?import\s+[\w.,\s]+$'
            for line in code.split('\n'):
                line = line.strip()
                if re.match(import_pattern, line):
                    analysis['imports'].append(line)
            
            # 提取函数定义
            func_pattern = r'def\s+(\w+)\s*\('
            for match in re.finditer(func_pattern, code):
                analysis['functions'].append(match.group(1))
            
            # 提取文件操作
            file_ops = ['open(', 'read_csv(', 'read_excel(', 'to_csv(', 'to_excel(', 'Document(', 'PdfReader(']
            for op in file_ops:
                if op in code:
                    analysis['file_operations'].append(op.replace('(', ''))
            
            # 提取路径变量
            path_pattern = r'(\w+_path)\s*='
            for match in re.finditer(path_pattern, code):
                analysis['variables'].append(match.group(1))
        
        except Exception as e:
            analysis['error'] = f"代码分析失败: {str(e)}"
        
        return analysis
    
    def _extract_json(self, text):
        """提取JSON内容"""
        try:
            return json.loads(text)
        except:
            if "```json" in text:
                text = text.split("```json")[1].split("```")[0].strip()
            elif "```" in text:
                text = text.split("```")[1].split("```")[0].strip()
            return json.loads(text)
    
    def _get_output_files(self, output_path):
        """获取输出文件列表"""
        if not output_path.exists():
            return "输出目录为空"
        
        files = []
        for file_path in output_path.rglob('*'):
            if file_path.is_file() and not file_path.name.startswith('.'):
                rel_path = file_path.relative_to(output_path)
                size = file_path.stat().st_size
                files.append(f"- {rel_path} ({self._format_size(size)})")
        
        if not files:
            return "输出目录为空"
        
        return "\n".join(files)
    
    def _format_size(self, bytes):
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if bytes < 1024.0:
                return f"{bytes:.1f}{unit}"
            bytes /= 1024.0
        return f"{bytes:.1f}TB"
    
    def _read_output_content(self, output_path, max_preview=500):
        """读取输出文件内容（精简版）"""
        if not output_path.exists():
            return "输出目录不存在"
        
        content_previews = []
        
        for file_path in output_path.rglob('*'):
            if not file_path.is_file() or file_path.name.startswith('.'):
                continue
            
            rel_path = file_path.relative_to(output_path)
            file_ext = file_path.suffix.lower()
            file_size = file_path.stat().st_size
            
            try:
                # CSV/Excel文件：只读取列名和行数
                if file_ext in ['.csv', '.xlsx', '.xls']:
                    import pandas as pd
                    try:
                        if file_ext == '.csv':
                            df = pd.read_csv(file_path, nrows=5)
                        else:
                            df = pd.read_excel(file_path, nrows=5)
                        
                        preview = f"\n【{rel_path}】\n"
                        preview += f"类型: {file_ext.upper()}, 大小: {self._format_size(file_size)}\n"
                        preview += f"列名: {', '.join(df.columns.tolist())}\n"
                        preview += f"数据行数: {len(df)}\n"
                        preview += f"前5行预览:\n{df.to_string(index=False, max_rows=5)}"
                        
                        content_previews.append(preview)
                    except Exception as e:
                        content_previews.append(f"\n【{rel_path}】\n读取失败: {str(e)}")
                
                # 文本文件：只读取前N个字符
                elif file_ext in ['.txt', '.json', '.xml', '.md', '.log']:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(max_preview)
                    
                    if len(content) >= max_preview:
                        content += "\n... (内容过长，已截断)"
                    
                    content_previews.append(f"\n【{rel_path}】\n{content}")
                
                # Word文件：只显示基本信息
                elif file_ext == '.docx':
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        content_previews.append(
                            f"\n【{rel_path}】\nWord文档，大小: {self._format_size(file_size)}, 段落数: {len(doc.paragraphs)}"
                        )
                    except:
                        content_previews.append(f"\n【{rel_path}】\nWord文档，大小: {self._format_size(file_size)}")
                
                # 其他文件：只显示基本信息
                else:
                    content_previews.append(
                        f"\n【{rel_path}】\n文件类型: {file_ext}, 大小: {self._format_size(file_size)}"
                    )
                    
            except Exception as e:
                content_previews.append(f"\n【{rel_path}】\n读取失败: {str(e)}")
        
        if not content_previews:
            return "输出目录为空"
        
        return "\n".join(content_previews)
